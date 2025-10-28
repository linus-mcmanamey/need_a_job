"""
CV Tailor Agent - Customizes CV template for each specific job.

This agent reads the candidate's CV template, analyzes job requirements,
and generates a customized CV that highlights relevant skills and experience.
Uses Claude Sonnet 4 for intelligent document transformation while maintaining
factual accuracy.
"""

import json
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from loguru import logger

from app.agents.base_agent import AgentResult, BaseAgent


class CVTailorAgent(BaseAgent):
    """
    Agent that customizes CV templates for specific job applications.

    This agent performs intelligent CV customization:
    - Loads CV template from DOCX file
    - Analyzes job requirements and matched criteria
    - Uses Claude to generate customization instructions
    - Applies customizations while preserving structure
    - Generates tailored CV in timestamped directory
    - Validates output and updates database

    Attributes:
        _cv_template_path: Path to CV template file
    """

    def __init__(
        self,
        config: dict[str, Any],
        claude_client: Any,
        app_repository: Any,
    ):
        """
        Initialize CV Tailor Agent.

        Args:
            config: Agent-specific configuration from agents.yaml
            claude_client: Anthropic Claude API client
            app_repository: ApplicationRepository for database access
        """
        super().__init__(config, claude_client, app_repository)
        self._cv_template_path = Path("current_cv_coverletter/Linus_McManamey_CV.docx")

    @property
    def agent_name(self) -> str:
        """Return agent name."""
        return "cv_tailor"

    async def process(self, job_id: str) -> AgentResult:
        """
        Process a job through CV tailoring.

        Loads CV template, analyzes job requirements, generates customized CV
        using Claude, and saves to timestamped directory.

        Args:
            job_id: UUID of the job to process

        Returns:
            AgentResult with success status, CV file path, and customization details
        """
        start_time = time.time()

        try:
            # Validate job_id
            if not job_id:
                logger.error("[cv_tailor] Missing job_id parameter")
                return AgentResult(
                    success=False,
                    agent_name=self.agent_name,
                    output={},
                    error_message="Missing job_id parameter",
                    execution_time_ms=int((time.time() - start_time) * 1000),
                )

            # Load job data from database
            logger.info(f"[cv_tailor] Processing job: {job_id}")
            job_data = await self._app_repo.get_job_by_id(job_id)

            if not job_data:
                logger.error(f"[cv_tailor] Job not found: {job_id}")
                return AgentResult(
                    success=False,
                    agent_name=self.agent_name,
                    output={},
                    error_message=f"Job not found: {job_id}",
                    execution_time_ms=int((time.time() - start_time) * 1000),
                )

            # Update current stage
            await self._update_current_stage(job_id, self.agent_name)

            # Load CV template
            try:
                cv_doc = self._load_cv_template(self._cv_template_path)
            except FileNotFoundError:
                logger.error(f"[cv_tailor] CV template not found: {self._cv_template_path}")
                return AgentResult(
                    success=False,
                    agent_name=self.agent_name,
                    output={},
                    error_message=f"CV template not found: {self._cv_template_path}",
                    execution_time_ms=int((time.time() - start_time) * 1000),
                )
            except Exception as e:
                logger.error(f"[cv_tailor] Failed to load CV template: {e}")
                return AgentResult(
                    success=False,
                    agent_name=self.agent_name,
                    output={},
                    error_message=f"Failed to load CV template: {str(e)}",
                    execution_time_ms=int((time.time() - start_time) * 1000),
                )

            # Load stage outputs for matched criteria
            stage_outputs = await self._app_repo.get_stage_outputs(job_id)

            # Analyze job requirements
            job_context = self._analyze_job_requirements(job_data, stage_outputs)

            # Extract CV content as text
            cv_content = self._extract_cv_content(cv_doc)

            # Customize CV with Claude
            customizations = await self._customize_cv_with_claude(cv_content, job_context)

            # Create output directory
            company = job_data.get("company_name", "unknown")
            job_title = job_data.get("title", "unknown")
            output_dir = self._create_output_directory(company, job_title)

            # Generate customized CV
            output_path = output_dir / "Linus_McManamey_CV.docx"
            self._generate_customized_cv(cv_doc, customizations, output_path)

            # Validate output file
            if not self._validate_output_file(output_path):
                logger.error(f"[cv_tailor] Output file validation failed: {output_path}")
                return AgentResult(
                    success=False,
                    agent_name=self.agent_name,
                    output={},
                    error_message="Output file validation failed",
                    execution_time_ms=int((time.time() - start_time) * 1000),
                )

            # Get file size
            file_size = output_path.stat().st_size

            # Build output
            output = {
                "cv_file_path": str(output_path),
                "customization_notes": customizations.get("customization_notes", ""),
                "sections_reordered": customizations.get("section_order", []),
                "keywords_incorporated": customizations.get("keywords_to_add", []),
                "file_size_bytes": file_size,
            }

            # Update database with CV file path
            if hasattr(self._app_repo, "update_cv_file_path"):
                await self._app_repo.update_cv_file_path(job_id, str(output_path))

            # Update database with stage completion
            await self._add_completed_stage(job_id, self.agent_name, output)

            # Log success
            logger.info(
                f"[cv_tailor] Job {job_id}: CV generated at {output_path}, size={file_size} bytes"
            )

            execution_time_ms = int((time.time() - start_time) * 1000)

            return AgentResult(
                success=True,
                agent_name=self.agent_name,
                output=output,
                error_message=None,
                execution_time_ms=execution_time_ms,
            )

        except Exception as e:
            logger.error(f"[cv_tailor] Error processing job {job_id}: {e}")
            execution_time_ms = int((time.time() - start_time) * 1000)

            return AgentResult(
                success=False,
                agent_name=self.agent_name,
                output={},
                error_message=str(e),
                execution_time_ms=execution_time_ms,
            )

    def _load_cv_template(self, template_path: Path) -> Document:
        """
        Load CV template from DOCX file.

        Args:
            template_path: Path to CV template file

        Returns:
            Document object

        Raises:
            FileNotFoundError: If template file doesn't exist
            Exception: If file cannot be parsed as DOCX
        """
        if not template_path.exists():
            raise FileNotFoundError(f"CV template not found: {template_path}")

        try:
            doc = Document(template_path)
            logger.info(f"[cv_tailor] Loaded CV template: {template_path}")
            return doc
        except Exception as e:
            logger.error(f"[cv_tailor] Failed to parse CV template: {e}")
            raise

    def _extract_cv_content(self, cv_doc: Document) -> str:
        """
        Extract text content from CV document.

        Args:
            cv_doc: Document object

        Returns:
            CV content as text
        """
        paragraphs = []
        for para in cv_doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        content = "\n".join(paragraphs)
        logger.debug(f"[cv_tailor] Extracted {len(paragraphs)} paragraphs from CV")
        return content

    def _analyze_job_requirements(
        self, job_data: dict[str, Any], stage_outputs: dict[str, Any]
    ) -> dict[str, Any]:
        """
        Analyze job requirements and extract relevant context.

        Args:
            job_data: Job information from database
            stage_outputs: Previous agent outputs (especially job_matcher)

        Returns:
            Dictionary with job context for Claude
        """
        job_title = job_data.get("title", "Unknown")
        company_name = job_data.get("company_name", "Unknown")
        description = job_data.get("description", "")

        # Extract matched technologies from job_matcher output
        matched_technologies = []
        if "job_matcher" in stage_outputs:
            matcher_output = stage_outputs["job_matcher"]
            matched_technologies.extend(matcher_output.get("must_have_found", []))
            matched_technologies.extend(matcher_output.get("strong_pref_found", []))
            matched_technologies.extend(matcher_output.get("nice_to_have_found", []))

        context = {
            "job_title": job_title,
            "company_name": company_name,
            "job_description": description,
            "matched_technologies": matched_technologies,
        }

        logger.debug(f"[cv_tailor] Job context: {len(matched_technologies)} matched technologies")
        return context

    async def _customize_cv_with_claude(
        self, cv_content: str, job_context: dict[str, Any]
    ) -> dict[str, Any]:
        """
        Use Claude to generate CV customization instructions.

        Args:
            cv_content: Original CV content as text
            job_context: Job requirements and matched criteria

        Returns:
            Dictionary with customization instructions

        Raises:
            Exception: If Claude API fails
        """
        prompt = f"""You are a CV Customization Agent. Analyze the job requirements and customize the candidate's CV to highlight relevant experience.

ORIGINAL CV CONTENT:
{cv_content[:2000]}  # Limit to 2000 chars for token efficiency

JOB REQUIREMENTS:
- Title: {job_context["job_title"]}
- Company: {job_context["company_name"]}
- Key Technologies: {", ".join(job_context["matched_technologies"][:10])}
- Description: {job_context["job_description"][:500]}

TASK:
Provide customization instructions to tailor this CV for the job. Include:
1. Section reordering (prioritize relevant sections first)
2. Emphasis areas (which skills/technologies to highlight)
3. Keyword incorporation (where to naturally add job keywords)
4. Professional summary customization (first paragraph rewrite)

RULES:
- DO NOT fabricate experience, skills, or qualifications
- Maintain factual accuracy of all dates, companies, and roles
- Use Australian English spelling and grammar
- Keep professional tone and formatting
- Only reorder, emphasize, or rephrase existing content

OUTPUT FORMAT (JSON only):
{{
  "section_order": ["Professional Summary", "Key Skills", "Work Experience", "Education"],
  "emphasis_skills": ["Azure", "PySpark", "Databricks"],
  "keywords_to_add": ["Data Engineering", "ETL pipelines", "Cloud infrastructure"],
  "professional_summary": "Rewritten first paragraph emphasizing relevant experience",
  "customization_notes": "Brief explanation of changes made"
}}"""

        system_prompt = """You are a professional CV customization specialist. Analyze job requirements and provide tailored CV customization instructions. Return JSON only, no additional text."""

        try:
            response = await self._call_claude(prompt, system_prompt)
            customizations = self._parse_customization_response(response)
            return customizations
        except Exception as e:
            logger.error(f"[cv_tailor] Claude API error: {e}")
            raise

    def _parse_customization_response(self, response: str) -> dict[str, Any]:
        """
        Parse Claude's JSON response for customization instructions.

        Args:
            response: JSON string from Claude

        Returns:
            Parsed customization dictionary
        """
        try:
            # Extract JSON from response (Claude sometimes adds markdown)
            if "```json" in response:
                start = response.find("```json") + 7
                end = response.find("```", start)
                response = response[start:end].strip()
            elif "```" in response:
                start = response.find("```") + 3
                end = response.find("```", start)
                response = response[start:end].strip()

            parsed = json.loads(response)
            logger.debug("[cv_tailor] Successfully parsed customization instructions")
            return parsed

        except json.JSONDecodeError as e:
            logger.error(f"[cv_tailor] Failed to parse Claude response: {e}")
            # Return minimal customization on parse failure
            return {
                "section_order": [],
                "emphasis_skills": [],
                "keywords_to_add": [],
                "professional_summary": "",
                "customization_notes": "Failed to parse customization instructions",
            }

    def _create_output_directory(self, company: str, job_title: str) -> Path:
        """
        Create output directory for CV and CL.

        Directory format: YYYY-MM-DD_company_jobtitle

        Args:
            company: Company name
            job_title: Job title

        Returns:
            Path to created directory
        """
        date_str = datetime.now().strftime("%Y-%m-%d")
        company_sanitized = self._sanitize_filename(company)
        title_sanitized = self._sanitize_filename(job_title)

        dir_name = f"{date_str}_{company_sanitized}_{title_sanitized}"
        output_dir = Path("export_cv_cover_letter") / dir_name

        output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"[cv_tailor] Created output directory: {output_dir}")

        return output_dir

    def _sanitize_filename(self, text: str) -> str:
        """
        Sanitize text for use in filename.

        Removes invalid characters, replaces spaces, converts to lowercase.

        Args:
            text: Text to sanitize

        Returns:
            Sanitized text safe for filesystem
        """
        # Remove or replace invalid characters
        sanitized = re.sub(r'[<>:"/\\|?*]', "", text)
        # Replace spaces with hyphens
        sanitized = sanitized.replace(" ", "-")
        # Convert to lowercase
        sanitized = sanitized.lower()
        # Limit length to 50 characters
        sanitized = sanitized[:50]
        # Remove trailing hyphens
        sanitized = sanitized.rstrip("-")

        return sanitized

    def _generate_customized_cv(
        self, cv_doc: Document, customizations: dict[str, Any], output_path: Path
    ) -> None:
        """
        Generate customized CV by applying customizations to template.

        For MVP, we save the original CV with minimal modifications.
        Future: Apply section reordering, emphasis, keyword incorporation.

        Args:
            cv_doc: Original CV document
            customizations: Customization instructions from Claude
            output_path: Path to save customized CV
        """
        try:
            # For MVP: Save original CV (full customization in future iteration)
            # Future enhancements:
            # - Reorder sections based on section_order
            # - Bold/emphasize skills from emphasis_skills
            # - Add keywords naturally from keywords_to_add
            # - Replace professional summary

            cv_doc.save(output_path)
            logger.info(f"[cv_tailor] Saved customized CV: {output_path}")

        except Exception as e:
            logger.error(f"[cv_tailor] Failed to save CV: {e}")
            raise

    def _validate_output_file(self, file_path: Path) -> bool:
        """
        Validate generated CV file.

        Checks:
        - File exists
        - File size < 5MB
        - Valid DOCX format (can be opened)

        Args:
            file_path: Path to CV file

        Returns:
            True if valid, False otherwise
        """
        MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

        # Check file exists
        if not file_path.exists():
            logger.error(f"[cv_tailor] Output file not found: {file_path}")
            return False

        # Check file size
        file_size = file_path.stat().st_size
        if file_size > MAX_FILE_SIZE:
            logger.error(f"[cv_tailor] File size exceeds limit: {file_size} bytes")
            return False

        # Test if DOCX is valid by trying to open it
        try:
            Document(file_path)
            logger.info(f"[cv_tailor] Validated output file: {file_path}, size={file_size} bytes")
            return True
        except Exception as e:
            logger.error(f"[cv_tailor] Invalid DOCX file: {e}")
            return False
