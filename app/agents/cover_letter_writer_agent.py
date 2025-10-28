"""
Cover Letter Writer Agent - Writes personalized cover letters for each job.

Uses Claude Sonnet 4 to generate tailored cover letters that address selection
criteria and demonstrate genuine interest. Extracts contact person name and
maintains professional Australian English tone.
"""

import re
import time
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from loguru import logger

from app.agents.base_agent import AgentResult, BaseAgent


class CoverLetterWriterAgent(BaseAgent):
    """
    Agent that writes personalized cover letters for job applications.

    Features:
    - Loads CL template from DOCX file
    - Extracts contact person name from job data
    - Uses Claude to generate personalized CL
    - Saves to same directory as CV
    - Validates output and updates database
    """

    def __init__(
        self,
        config: dict[str, Any],
        claude_client: Any,
        app_repository: Any,
    ):
        """Initialize Cover Letter Writer Agent."""
        super().__init__(config, claude_client, app_repository)
        self._cl_template_path = Path("current_cv_coverletter/Linus_McManamey_CL.docx")

    @property
    def agent_name(self) -> str:
        """Return agent name."""
        return "cover_letter_writer"

    async def process(self, job_id: str) -> AgentResult:
        """
        Process a job through cover letter writing.

        Args:
            job_id: UUID of the job to process

        Returns:
            AgentResult with success status, CL file path, contact person
        """
        start_time = time.time()

        try:
            # Validate job_id
            if not job_id:
                logger.error("[cover_letter_writer] Missing job_id parameter")
                return AgentResult(
                    success=False,
                    agent_name=self.agent_name,
                    output={},
                    error_message="Missing job_id parameter",
                    execution_time_ms=int((time.time() - start_time) * 1000),
                )

            # Load job data
            logger.info(f"[cover_letter_writer] Processing job: {job_id}")
            job_data = await self._app_repo.get_job_by_id(job_id)

            if not job_data:
                logger.error(f"[cover_letter_writer] Job not found: {job_id}")
                return AgentResult(
                    success=False,
                    agent_name=self.agent_name,
                    output={},
                    error_message=f"Job not found: {job_id}",
                    execution_time_ms=int((time.time() - start_time) * 1000),
                )

            # Update current stage
            await self._update_current_stage(job_id, self.agent_name)

            # Load CL template (validates template exists and is valid DOCX)
            try:
                self._load_cl_template(self._cl_template_path)
            except FileNotFoundError:
                logger.error(f"[cover_letter_writer] Template not found: {self._cl_template_path}")
                return AgentResult(
                    success=False,
                    agent_name=self.agent_name,
                    output={},
                    error_message=f"CL template not found: {self._cl_template_path}",
                    execution_time_ms=int((time.time() - start_time) * 1000),
                )

            # Load stage outputs
            stage_outputs = await self._app_repo.get_stage_outputs(job_id)

            # Extract contact person
            contact_person, extraction_method = self._extract_contact_person(job_data)

            # Prepare job context
            job_context = self._prepare_job_context(job_data, stage_outputs)

            # Generate cover letter with Claude
            cl_text = await self._generate_cover_letter_with_claude(job_context, contact_person)

            # Get output directory from cv_tailor stage
            cv_file_path = stage_outputs.get("cv_tailor", {}).get("cv_file_path", "")
            if cv_file_path:
                output_dir = Path(cv_file_path).parent
            else:
                # Fallback: create new directory with sanitized company name
                from datetime import datetime

                date_str = datetime.now().strftime("%Y-%m-%d")
                company = self._sanitize_filename(job_data.get("company_name", "unknown"))
                output_dir = Path(f"export_cv_cover_letter/{date_str}_{company}")
                output_dir.mkdir(parents=True, exist_ok=True)

            # Save cover letter
            output_path = output_dir / "Linus_McManamey_CL.docx"
            self._create_cover_letter_docx(cl_text, output_path)

            # Validate output
            if not self._validate_output_file(output_path):
                logger.error(f"[cover_letter_writer] Validation failed: {output_path}")
                return AgentResult(
                    success=False,
                    agent_name=self.agent_name,
                    output={},
                    error_message="Output file validation failed",
                    execution_time_ms=int((time.time() - start_time) * 1000),
                )

            # Get file size
            file_size = output_path.stat().st_size

            # Build output
            output = {
                "cl_file_path": str(output_path),
                "contact_person_name": contact_person,
                "extraction_method": extraction_method,
                "file_size_bytes": file_size,
            }

            # Update database
            if hasattr(self._app_repo, "update_cl_file_path"):
                await self._app_repo.update_cl_file_path(job_id, str(output_path))
            if hasattr(self._app_repo, "update_contact_person"):
                await self._app_repo.update_contact_person(job_id, contact_person)

            await self._add_completed_stage(job_id, self.agent_name, output)

            logger.info(
                f"[cover_letter_writer] Job {job_id}: CL generated, "
                f"contact={contact_person}, size={file_size}"
            )

            execution_time_ms = int((time.time() - start_time) * 1000)

            return AgentResult(
                success=True,
                agent_name=self.agent_name,
                output=output,
                error_message=None,
                execution_time_ms=execution_time_ms,
            )

        except Exception as e:
            logger.error(f"[cover_letter_writer] Error processing job {job_id}: {e}")
            execution_time_ms = int((time.time() - start_time) * 1000)

            return AgentResult(
                success=False,
                agent_name=self.agent_name,
                output={},
                error_message=str(e),
                execution_time_ms=execution_time_ms,
            )

    def _load_cl_template(self, template_path: Path) -> Document:
        """Load CL template from DOCX file."""
        if not template_path.exists():
            raise FileNotFoundError(f"CL template not found: {template_path}")

        try:
            doc = Document(template_path)
            logger.info(f"[cover_letter_writer] Loaded template: {template_path}")
            return doc
        except Exception as e:
            logger.error(f"[cover_letter_writer] Failed to parse template: {e}")
            raise

    def _extract_contact_person(self, job_data: dict[str, Any]) -> tuple[str, str]:
        """
        Extract contact person name from job data.

        Returns:
            Tuple of (name, extraction_method)
        """
        # Try email first
        email = job_data.get("contact_email", "")
        if email:
            name = self._parse_name_from_email(email)
            if name and name != "Unknown":
                logger.debug(f"[cover_letter_writer] Extracted name from email: {name}")
                return (name, "email")

        # Default
        logger.debug("[cover_letter_writer] Using default contact name")
        return ("Hiring Manager", "default")

    def _parse_name_from_email(self, email: str) -> str:
        """
        Parse name from email address.

        Examples:
            jane.smith@acme.com -> Jane Smith
            jsmith@acme.com -> J Smith
        """
        try:
            local_part = email.split("@")[0]
            # Handle firstname.lastname
            if "." in local_part:
                parts = local_part.split(".")
                return " ".join(p.capitalize() for p in parts)
            # Handle firstnamelastname or single name
            else:
                # Just capitalize first letter
                return local_part[0].upper() + " " + local_part[1:].capitalize()
        except Exception:
            return "Unknown"

    def _sanitize_filename(self, text: str) -> str:
        """
        Sanitize text for use in filename to prevent path traversal.

        Removes invalid characters and path separators for cross-platform safety.

        Args:
            text: Text to sanitize

        Returns:
            Sanitized string safe for use in filename
        """
        # Remove or replace invalid filesystem characters and path separators
        sanitized = re.sub(r'[<>:"/\\|?*]', "", text)
        # Replace spaces with hyphens
        sanitized = sanitized.replace(" ", "-")
        # Convert to lowercase
        sanitized = sanitized.lower()
        # Limit length to 50 characters
        sanitized = sanitized[:50]
        # Remove trailing hyphens
        sanitized = sanitized.rstrip("-")
        return sanitized if sanitized else "unknown"

    def _prepare_job_context(
        self, job_data: dict[str, Any], stage_outputs: dict[str, Any]
    ) -> dict[str, Any]:
        """Prepare job context for Claude."""
        # Extract matched technologies
        matched_techs = []
        if "job_matcher" in stage_outputs:
            matcher_output = stage_outputs["job_matcher"]
            matched_techs.extend(matcher_output.get("must_have_found", []))
            matched_techs.extend(matcher_output.get("strong_pref_found", []))

        context = {
            "company_name": job_data.get("company_name", "Unknown"),
            "job_title": job_data.get("title", "Unknown"),
            "job_description": job_data.get("description", ""),
            "matched_technologies": matched_techs,
        }

        return context

    async def _generate_cover_letter_with_claude(
        self, job_context: dict[str, Any], contact_person: str
    ) -> str:
        """Generate cover letter using Claude."""
        prompt = f"""You are a professional Cover Letter Writer. Write a personalized, compelling cover letter for this job application.

JOB DETAILS:
- Company: {job_context['company_name']}
- Title: {job_context['job_title']}
- Description: {job_context['job_description'][:800]}

MATCHED CRITERIA:
{', '.join(job_context['matched_technologies'][:10])}

CONTACT PERSON:
{contact_person}

TASK:
Write a complete, professional cover letter that:
1. Opens with company-specific introduction addressing {contact_person}
2. Addresses 2-3 key selection criteria from the job
3. Highlights relevant technical skills and achievements
4. Shows genuine interest in the role and company
5. Includes clear call to action
6. Uses professional but friendly tone
7. Australian English spelling and grammar
8. 3-4 paragraphs, ~300-400 words

IMPORTANT:
- Start with "Dear {contact_person},"
- Be specific to this company and role
- Show enthusiasm without being over-the-top
- End with professional closing

OUTPUT FORMAT (plain text only, no JSON):
[Full cover letter text]"""

        system_prompt = """You are an expert cover letter writer. Write compelling, personalized cover letters that demonstrate fit and genuine interest. Use Australian English. Output plain text only."""

        try:
            response = await self._call_claude(prompt, system_prompt)
            logger.debug(f"[cover_letter_writer] Generated CL ({len(response)} chars)")
            return response
        except Exception as e:
            logger.error(f"[cover_letter_writer] Claude API error: {e}")
            raise

    def _create_cover_letter_docx(self, cl_text: str, output_path: Path) -> None:
        """Create cover letter DOCX from text."""
        try:
            doc = Document()

            # Add paragraphs from text
            for paragraph_text in cl_text.split("\n\n"):
                if paragraph_text.strip():
                    p = doc.add_paragraph(paragraph_text.strip())
                    # Set font
                    for run in p.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(11)

            doc.save(output_path)
            logger.info(f"[cover_letter_writer] Saved CL: {output_path}")

        except Exception as e:
            logger.error(f"[cover_letter_writer] Failed to create CL DOCX: {e}")
            raise

    def _validate_output_file(self, file_path: Path) -> bool:
        """Validate generated CL file."""
        MAX_FILE_SIZE = 2 * 1024 * 1024  # 2MB

        if not file_path.exists():
            logger.error(f"[cover_letter_writer] File not found: {file_path}")
            return False

        file_size = file_path.stat().st_size
        if file_size > MAX_FILE_SIZE:
            logger.error(f"[cover_letter_writer] File too large: {file_size}")
            return False

        try:
            Document(file_path)
            logger.info(f"[cover_letter_writer] Validated: {file_path}")
            return True
        except Exception as e:
            logger.error(f"[cover_letter_writer] Invalid DOCX: {e}")
            return False
